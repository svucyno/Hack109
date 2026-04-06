import re
import uuid
from io import BytesIO
from datetime import datetime
from urllib.error import HTTPError, URLError
from urllib.parse import urlparse
from urllib.request import Request, urlopen
from django.utils import timezone

from rest_framework import status
from rest_framework.permissions import AllowAny
from rest_framework.response import Response
from rest_framework.views import APIView

from .models import ResumeUploadRecord
from .storage import (
    generate_presigned_upload_url,
    get_s3_object_metadata,
    is_s3_configured,
    read_resume_file,
    store_resume_file,
)

try:
    from pypdf import PdfReader
except ImportError:  # pragma: no cover - optional dependency
    PdfReader = None

try:
    from docx import Document as DocxDocument
except ImportError:  # pragma: no cover - optional dependency
    DocxDocument = None


UPLOAD_STORE: dict[str, dict] = {}
PARSE_STORE: dict[str, dict] = {}

SUPPORTED_EXTENSIONS = ('.pdf', '.docx')
DEFAULT_SKILLS = [
    'Python',
    'Django',
    'REST',
    'SQL',
    'Docker',
    'Kubernetes',
    'React',
    'Node.js',
]
SKILL_ALIASES: dict[str, list[str]] = {
    'Python': ['python'],
    'Django': ['django'],
    'REST': ['rest', 'rest api', 'restful'],
    'SQL': ['sql', 'mysql', 'postgresql', 'postgres'],
    'Docker': ['docker'],
    'Kubernetes': ['kubernetes', 'k8s'],
    'React': ['react', 'reactjs', 'react.js'],
    'Node.js': ['node', 'node.js', 'nodejs'],
    'Git': ['git', 'github', 'gitlab'],
    'AWS': ['aws', 'amazon web services'],
    'Azure': ['azure', 'microsoft azure'],
    'CI/CD': ['ci/cd', 'cicd', 'jenkins', 'github actions'],
    'Linux': ['linux', 'ubuntu'],
    'Machine Learning': ['machine learning', 'ml'],
    'TensorFlow': ['tensorflow'],
    'Pandas': ['pandas'],
    'NumPy': ['numpy'],
}
ROLE_SKILL_MAP: dict[str, list[str]] = {
    'software engineer': ['Python', 'SQL', 'Git', 'REST'],
    'backend engineer': ['Python', 'Django', 'REST', 'SQL', 'Docker'],
    'backend developer': ['Python', 'Django', 'REST', 'SQL', 'Docker'],
    'full stack developer': ['React', 'Node.js', 'SQL', 'REST', 'Git'],
    'full-stack developer': ['React', 'Node.js', 'SQL', 'REST', 'Git'],
    'devops engineer': ['Docker', 'Kubernetes', 'Linux', 'CI/CD', 'AWS', 'Azure'],
    'cloud engineer': ['AWS', 'Azure', 'Docker', 'Kubernetes', 'Linux'],
    'data engineer': ['Python', 'SQL', 'Pandas', 'NumPy'],
    'machine learning engineer': ['Python', 'Machine Learning', 'TensorFlow', 'Pandas', 'NumPy'],
}
ROLE_KEYWORDS = [
    'software engineer',
    'backend developer',
    'backend engineer',
    'full stack developer',
    'full-stack developer',
    'data engineer',
    'data analyst',
    'machine learning engineer',
    'devops engineer',
    'cloud engineer',
    'python developer',
]
SECTION_HEADERS = {
    'experience': ['experience', 'work experience', 'professional experience', 'employment history'],
    'projects': ['projects', 'project'],
    'education': ['education', 'academic background', 'academics'],
    'skills': ['skills', 'technical skills', 'core skills'],
}
ROLE_NOISE_TERMS = {'project', 'projects', 'education', 'skills', 'summary', 'profile'}
ROLE_VERB_NOISE = {'deployed', 'managed', 'configured', 'implemented', 'designed', 'developed'}
ROLE_PREFIX_VERB_NOISE = {'engage', 'engaged', 'worked', 'working', 'collaborated', 'supporting', 'learning'}
DATE_RANGE_RE = re.compile(
    r'(?P<start>(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)?\s*\d{4})\s*[-–to]+\s*'
    r'(?P<end>present|current|now|(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)?\s*\d{4})',
    re.IGNORECASE,
)
TITLE_RE = re.compile(
    r'\b([A-Za-z][A-Za-z/&\- ]{2,40}(Engineer|Developer|Lead|Analyst|Architect|Intern|Manager))\b',
    re.IGNORECASE,
)
TITLE_FALLBACK_RE = re.compile(
    r'\b([A-Za-z&/\-]{2,20}(?:\s+[A-Za-z&/\-]{1,20}){0,3}\s(?:Engineer|Developer|Lead|Analyst|Architect|Intern|Manager))\b',
    re.IGNORECASE,
)
YEAR_RANGE_RE = re.compile(
    r'\b((?:19|20)\d{2})\s*(?:-|–|to)\s*((?:19|20)\d{2}|present|current|now)\b',
    re.IGNORECASE,
)
INSTITUTION_RE = re.compile(
    r'([A-Z][A-Za-z0-9&(),.\- ]{3,120}(?:University|Institute|College|School|Academy|Technological University))',
    re.IGNORECASE,
)
DEGREE_PATTERNS = [
    re.compile(r'\bBachelor\s+of\s+[A-Za-z &]+', re.IGNORECASE),
    re.compile(r'\bMaster\s+of\s+[A-Za-z &]+', re.IGNORECASE),
    re.compile(r'\bB\.?\s?Tech\b(?:\s+in\s+[A-Za-z &]+)?', re.IGNORECASE),
    re.compile(r'\bB\.?\s?E\b(?:\s+in\s+[A-Za-z &]+)?', re.IGNORECASE),
    re.compile(r'\bM\.?\s?Tech\b(?:\s+in\s+[A-Za-z &]+)?', re.IGNORECASE),
    re.compile(r'\bM\.?\s?E\b(?:\s+in\s+[A-Za-z &]+)?', re.IGNORECASE),
    re.compile(r'\bPh\.?D\b(?:\s+in\s+[A-Za-z &]+)?', re.IGNORECASE),
]
EDUCATION_KEYWORDS = [
    'b.tech',
    'b.e',
    'bachelor',
    'm.tech',
    'm.e',
    'master',
    'phd',
    'university',
    'college',
]
URL_RE = re.compile(r'(?i)\b(?:https?://|www\.)[^\s<>()]+')
HREF_URL_RE = re.compile(r'(?i)href\s*=\s*[\"\']([^\"\']+)[\"\']')
MARKDOWN_URL_RE = re.compile(r'\[[^\]]+\]\(([^)]+)\)')
GITHUB_HINT = 'github.com/'
NON_PROD_HOST_BLOCKLIST = {
    'linkedin.com',
    'gitlab.com',
    'bitbucket.org',
    'leetcode.com',
    'hackerrank.com',
    'behance.net',
    'dribbble.com',
    'medium.com',
    'twitter.com',
    'x.com',
    'facebook.com',
    'instagram.com',
}
MAX_LINKS_TO_VERIFY = 8
URL_VERIFY_TIMEOUT_SECONDS = 4


def _valid_resume_filename(filename: str) -> bool:
    return bool(filename) and filename.lower().endswith(SUPPORTED_EXTENSIONS)


def _unsupported_file_detail(filename: str) -> str:
    lowered = (filename or '').lower()
    if lowered.endswith('.doc'):
        return 'Unsupported file type. Legacy .doc is not supported; please convert it to .docx and upload again.'
    return 'Unsupported file type. Only PDF and DOCX are accepted.'


def _extract_text(content: bytes) -> str:
    def _repair_unicode(value: str) -> str:
        repaired = value.replace('\u00a0', ' ')
        replacements = {
            '•': '-',
            '–': '-',
            '—': '-',
            'â\x80\x93': '-',
            'â\x80\x94': '-',
            'â\x80\x98': "'",
            'â\x80\x99': "'",
            'â\x80\x9c': '"',
            'â\x80\x9d': '"',
            'â\x80\xa2': '-',
            'â\x80\xa6': '...',
            'â¢': '-',
            'â€“': '-',
            'â€”': '-',
            'Ã©': 'e',
            'Ã¨': 'e',
            'Â': '',
        }
        for bad, good in replacements.items():
            repaired = repaired.replace(bad, good)

        if any(marker in repaired for marker in ('Ã', 'â', 'Â')):
            try:
                maybe_fixed = repaired.encode('latin1', errors='ignore').decode('utf-8', errors='ignore')
                if maybe_fixed and maybe_fixed.count('Ã') + maybe_fixed.count('â') < repaired.count('Ã') + repaired.count('â'):
                    repaired = maybe_fixed
            except Exception:
                pass

        repaired = re.sub(r'\sâ\s', ' - ', repaired)
        repaired = re.sub(r'â(?=\s|\W|$)', '-', repaired)
        repaired = repaired.replace('â', '-')
        return repaired

    def _normalize(value: str) -> str:
        cleaned = _repair_unicode(value)
        lines = [re.sub(r'\s+', ' ', line).strip() for line in cleaned.splitlines()]
        normalized = '\n'.join([line for line in lines if line])
        return normalized[:15000]

    if content.startswith(b'%PDF') and PdfReader is not None:
        try:
            reader = PdfReader(BytesIO(content))
            pages = [page.extract_text() or '' for page in reader.pages]
            extracted = _normalize('\n'.join(pages).strip())
            if extracted:
                return extracted
        except Exception:
            pass

    # DOCX files are zip containers (PK header). Parse paragraph/cell text when available.
    if content.startswith(b'PK') and DocxDocument is not None:
        try:
            document = DocxDocument(BytesIO(content))
            lines: list[str] = []
            lines.extend([p.text for p in document.paragraphs if p.text and p.text.strip()])
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            lines.append(cell_text)
            extracted = _normalize('\n'.join(lines).strip())
            if extracted:
                return extracted
        except Exception:
            pass

    extracted = _normalize(content.decode('utf-8', errors='ignore').strip())
    if not extracted:
        return 'Scanned or binary resume uploaded. OCR fallback scheduled in async pipeline.'
    return extracted


def _infer_skills(text: str) -> list[str]:
    lowered = text.lower()
    found: list[str] = []

    for canonical, aliases in SKILL_ALIASES.items():
        if any(re.search(rf'\b{re.escape(alias)}\b', lowered) for alias in aliases):
            found.append(canonical)

    # Preserve original base-skill behavior as fallback for compatibility.
    for skill in DEFAULT_SKILLS:
        if skill not in found and skill.lower() in lowered:
            found.append(skill)

    return found[:20]


def _infer_job_relevant_skills(roles: list[str], skills: list[str]) -> dict:
    role_requirements: dict[str, list[str]] = {}
    required_union: set[str] = set()
    skills_set = set(skills)

    for role in roles:
        key = role.lower()
        requirements = ROLE_SKILL_MAP.get(key)
        if not requirements:
            for mapped_role, mapped_requirements in ROLE_SKILL_MAP.items():
                if mapped_role in key:
                    requirements = mapped_requirements
                    break
        if requirements:
            role_requirements[role] = requirements
            required_union.update(requirements)

    matched = [skill for skill in sorted(required_union) if skill in skills_set]
    missing = [skill for skill in sorted(required_union) if skill not in skills_set]

    return {
        'role_requirements': role_requirements,
        'matched': matched,
        'missing': missing,
    }


def _clean_lines(text: str) -> list[str]:
    return [line.strip() for line in text.splitlines() if line.strip()]


def _normalize_line(line: str) -> str:
    normalized = re.sub(r'\s+', ' ', line).strip(' |-:;,.')
    return normalized


def _shorten_line(line: str, max_len: int = 160) -> str:
    normalized = _normalize_line(line).lstrip('-').strip()
    if len(normalized) <= max_len:
        return normalized
    sentence_cut = normalized.find('.')
    if 0 < sentence_cut <= max_len:
        return normalized[:sentence_cut].strip()
    return normalized[:max_len].rstrip()


def _extract_sections(text: str) -> dict[str, list[str]]:
    lines = _clean_lines(text)
    sections: dict[str, list[str]] = {'experience': [], 'projects': [], 'education': [], 'skills': []}
    current_section: str | None = None

    for raw_line in lines:
        line = _normalize_line(raw_line)
        line_lower = line.lower()

        matched_header = None
        for section, aliases in SECTION_HEADERS.items():
            if line_lower in aliases:
                matched_header = section
                break

        if matched_header:
            current_section = matched_header
            continue

        if current_section:
            sections[current_section].append(line)

    return sections


def _parse_year_token(token: str) -> int | None:
    match = re.search(r'(19\d{2}|20\d{2})', token)
    if not match:
        return None
    return int(match.group(1))


def _infer_roles(text: str) -> list[str]:
    lowered = text.lower()
    sections = _extract_sections(text)
    found: list[str] = []

    for role in ROLE_KEYWORDS:
        if role in lowered:
            found.append(role.title())

    role_source_lines = sections['experience'] or _clean_lines(text)
    candidate_tokens: list[str] = []
    for line in role_source_lines:
        candidate_tokens.extend(re.split(r'[|\n;]', line))
    if not candidate_tokens:
        candidate_tokens = re.split(r'[|\n;]', text)

    for token in candidate_tokens:
        line = _normalize_line(token)
        if not line or len(line) > 140:
            continue
        line_lower = line.lower()
        if any(noise in line_lower for noise in ROLE_NOISE_TERMS):
            continue
        if any(verb in line_lower for verb in ROLE_VERB_NOISE):
            continue
        first_word = line_lower.split(' ', 1)[0]
        if first_word in ROLE_PREFIX_VERB_NOISE:
            continue
        if 'http' in line_lower or 'github' in line_lower or 'linkedin' in line_lower or '@' in line_lower:
            continue
        for match in TITLE_RE.findall(line):
            role = _normalize_line(match[0]).title()
            role_first_word = role.lower().split(' ', 1)[0]
            if role_first_word in ROLE_PREFIX_VERB_NOISE:
                continue
            if ' With ' in role:
                continue
            if 2 <= len(role.split()) <= 6 and role not in found:
                found.append(role)

    if not found:
        for match in TITLE_FALLBACK_RE.findall(text):
            role = _normalize_line(match)
            role_lower = role.lower()
            if any(noise in role_lower for noise in ROLE_NOISE_TERMS):
                continue
            if any(verb in role_lower for verb in ROLE_VERB_NOISE):
                continue
            first_word = role_lower.split(' ', 1)[0]
            if first_word in ROLE_PREFIX_VERB_NOISE:
                continue
            if ' with ' in role_lower:
                continue
            if 2 <= len(role.split()) <= 5:
                role_title = role.title()
                if role_title not in found:
                    found.append(role_title)

    return found[:5]


def _infer_projects(text: str) -> list[str]:
    sections = _extract_sections(text)
    lines = sections['projects'] or _clean_lines(text)
    project_lines = []
    for line in lines:
        line_lower = line.lower()
        if 'project' in line_lower or 'http' in line_lower:
            cleaned = line
            if ':' in cleaned and 'project' in cleaned.lower():
                cleaned = cleaned.split(':', 1)[1]
            project_lines.append(_shorten_line(cleaned))
    if project_lines:
        return [line for line in project_lines[:5] if line]

    # Fallback: pick actionable lines that look like work highlights.
    highlights = [
        line for line in lines
        if any(keyword in line.lower() for keyword in ['built', 'developed', 'implemented', 'designed'])
    ]
    return [_shorten_line(line) for line in highlights[:5]]


def _extract_education_entry(value: str) -> dict[str, str]:
    line = _shorten_line(value, max_len=220)

    year_range = ''
    year_match = YEAR_RANGE_RE.search(line)
    if year_match:
        year_range = f"{year_match.group(1)}-{year_match.group(2).title()}"

    degree = ''
    for pattern in DEGREE_PATTERNS:
        match = pattern.search(line)
        if match:
            degree = _normalize_line(match.group(0)).title()
            break

    institution = ''
    inst_match = INSTITUTION_RE.search(line)
    if inst_match:
        institution = _normalize_line(inst_match.group(1))

    if not institution:
        tokens = re.split(r'\||,', line)
        for token in tokens:
            candidate = _normalize_line(token)
            if any(word in candidate.lower() for word in ['university', 'institute', 'college', 'school', 'academy']):
                institution = candidate
                break

    if not degree:
        for token in re.split(r'\||,', line):
            candidate = _normalize_line(token)
            lower = candidate.lower()
            if any(word in lower for word in ['b.tech', 'bachelor', 'master', 'm.tech', 'phd', 'b.e', 'm.e']):
                degree = candidate
                break

    return {
        'institution': institution,
        'degree': degree,
        'year_range': year_range,
    }


def _infer_education(text: str) -> list[dict[str, str]]:
    sections = _extract_sections(text)
    lines = sections['education'] or _clean_lines(text)
    entries: list[dict[str, str]] = []
    for line in lines:
        line_lower = line.lower()
        if any(keyword in line_lower for keyword in EDUCATION_KEYWORDS):
            cleaned = line
            if 'education' in line_lower:
                idx = line_lower.find('education')
                cleaned = line[idx + len('education'):]
            parsed = _extract_education_entry(cleaned)
            if parsed['institution'] or parsed['degree'] or parsed['year_range']:
                entries.append(parsed)

    # Guarantee at least one structured education object when resume has education-related content.
    if not entries and any(keyword in text.lower() for keyword in EDUCATION_KEYWORDS):
        entries.append({'institution': '', 'degree': '', 'year_range': ''})

    return entries[:5]


def _infer_years_experience(text: str) -> int | None:
    sections = _extract_sections(text)
    experience_text = '\n'.join(sections['experience']) or text

    patterns = [
        r'(\d{1,2})\+?\s+years?',
        r'(\d{1,2})\+?\s*yrs?',
        r'experience\s+of\s+(\d{1,2})\s+years?',
        r'experience\s*[:\-]?\s*(\d{1,2})',
    ]
    for pattern in patterns:
        match = re.search(pattern, experience_text, flags=re.IGNORECASE)
        if match:
            return int(match.group(1))

    current_year = datetime.now().year
    starts: list[int] = []
    ends: list[int] = []
    for match in DATE_RANGE_RE.finditer(experience_text):
        start_year = _parse_year_token(match.group('start'))
        end_token = match.group('end').lower()
        end_year = current_year if end_token in {'present', 'current', 'now'} else _parse_year_token(end_token)
        if start_year and end_year and end_year >= start_year:
            starts.append(start_year)
            ends.append(end_year)

    if starts and ends:
        return max(0, max(ends) - min(starts))
    return None


def _strip_url_trailing_punctuation(url: str) -> str:
    return url.rstrip('.,;:!?)]}\'\"')


def _extract_links(text: str) -> list[str]:
    links: list[str] = []
    seen: set[str] = set()

    def _add(candidate: str) -> None:
        raw = _strip_url_trailing_punctuation(candidate.strip())
        if not raw:
            return
        lowered = raw.lower()
        normalized = raw

        if lowered.startswith('www.') or any(lowered.startswith(hint) for hint in PROFILE_URL_HINTS):
            normalized = f'https://{raw}'

        parsed = urlparse(normalized)
        if parsed.scheme not in {'http', 'https'} or not parsed.netloc:
            return
        if not _is_allowed_resume_link(normalized):
            return

        dedupe_key = normalized.lower()
        if dedupe_key in seen:
            return
        seen.add(dedupe_key)
        links.append(normalized)

    for match in URL_RE.findall(text):
        _add(match)

    for href_url in HREF_URL_RE.findall(text):
        _add(href_url)

    for embedded_url in MARKDOWN_URL_RE.findall(text):
        _add(embedded_url)

    lowered = text.lower()
    start = 0
    while True:
        idx = lowered.find(GITHUB_HINT, start)
        if idx == -1:
            break
        end = idx
        while end < len(text) and not text[end].isspace():
            end += 1
        _add(text[idx:end])
        start = end

    return links[:20]


def _is_allowed_resume_link(url: str) -> bool:
    parsed = urlparse(url)
    host = parsed.netloc.lower()
    if host.startswith('www.'):
        host = host[4:]

    if not host or host in {'localhost', '127.0.0.1'}:
        return False

    if host == 'github.com' or host.endswith('.github.com'):
        return True

    if host in NON_PROD_HOST_BLOCKLIST:
        return False

    # Treat non-blocklisted public domains as production/portfolio links.
    return '.' in host


def _classify_link(url: str) -> str:
    parsed = urlparse(url)
    host = parsed.netloc.lower()
    path_parts = [part for part in parsed.path.split('/') if part]

    if 'github.com' in host:
        if len(path_parts) == 1:
            return 'github_profile'
        if len(path_parts) >= 2:
            return 'github_repository'
        return 'github'
    return 'production_link'


def _verify_single_link(url: str) -> dict:
    parsed = urlparse(url)
    result = {
        'url': url,
        'domain': parsed.netloc.lower(),
        'type': _classify_link(url),
        'reachable': False,
        'status_code': None,
        'verified_at': timezone.now().isoformat(),
        'error': '',
    }

    if parsed.scheme not in {'http', 'https'} or not parsed.netloc:
        result['error'] = 'Invalid URL format'
        return result

    headers = {'User-Agent': 'GetHired-LinkVerifier/1.0'}
    methods = ['HEAD', 'GET']
    for method in methods:
        try:
            request = Request(url, headers=headers, method=method)
            with urlopen(request, timeout=URL_VERIFY_TIMEOUT_SECONDS) as response:
                code = int(getattr(response, 'status', 0) or 0)
                result['status_code'] = code
                result['reachable'] = 200 <= code < 400
                result['error'] = '' if result['reachable'] else f'HTTP {code}'
                return result
        except HTTPError as exc:
            code = int(exc.code)
            result['status_code'] = code
            result['reachable'] = 200 <= code < 400
            result['error'] = '' if result['reachable'] else f'HTTP {code}'
            if method == 'GET' or code in {401, 403, 404, 410}:
                return result
        except URLError as exc:
            result['error'] = str(getattr(exc, 'reason', exc))
            if method == 'GET':
                return result
        except Exception as exc:
            result['error'] = str(exc)
            if method == 'GET':
                return result

    return result


def _verify_links(urls: list[str]) -> list[dict]:
    return [_verify_single_link(url) for url in urls[:MAX_LINKS_TO_VERIFY]]


def _build_parsed_payload(reference_no: str, text: str, storage_backend: str, storage_key: str) -> dict:
    skills = _infer_skills(text)
    roles = _infer_roles(text)
    job_relevant = _infer_job_relevant_skills(roles, skills)
    projects = _infer_projects(text)
    education = _infer_education(text)
    years = _infer_years_experience(text)
    links = _extract_links(text)
    verified_links = _verify_links(links)

    return {
        'reference_no': reference_no,
        'structured_profile': {
            'skills': skills,
            'roles': roles,
            'job_relevant_skills': job_relevant,
            'years_experience': years,
            'projects': projects,
            'education': education,
            'links': links,
            'verified_links': verified_links,
        },
        'parse_meta': {
            'source': 'phase2-light-parser',
            'ocr_fallback_used': 'ocr fallback scheduled' in text.lower(),
            'storage_backend': storage_backend,
            'storage_key': storage_key,
            'verified_links_count': len(verified_links),
        },
    }


def _parse_and_persist_record(record: ResumeUploadRecord) -> dict:
    content = read_resume_file(record.storage_backend, record.storage_key)
    text = _extract_text(content)

    UPLOAD_STORE[record.reference_no] = {
        'filename': record.original_filename,
        'content_type': record.content_type,
        'storage_backend': record.storage_backend,
        'storage_key': record.storage_key,
        'text_preview': text,
    }

    parsed_payload = _build_parsed_payload(
        reference_no=record.reference_no,
        text=text,
        storage_backend=record.storage_backend,
        storage_key=record.storage_key,
    )

    PARSE_STORE[record.reference_no] = parsed_payload
    record.status = 'parsed'
    record.parsed_json = parsed_payload
    record.parsed_at = timezone.now()
    record.save(update_fields=['status', 'parsed_json', 'parsed_at', 'updated_at'])
    return parsed_payload


class ResumeUploadView(APIView):
    permission_classes = [AllowAny]

    def post(self, request):
        uploaded = request.FILES.get('resume')
        if not uploaded:
            return Response(
                {'detail': "File field 'resume' is required."},
                status=status.HTTP_400_BAD_REQUEST,
            )

        filename = uploaded.name or ''
        if not _valid_resume_filename(filename):
            return Response(
                {'detail': _unsupported_file_detail(filename)},
                status=status.HTTP_400_BAD_REQUEST,
            )

        reference_no = f"REF-{str(uuid.uuid4())[:8].upper()}"
        file_content = uploaded.read()
        text_preview = _extract_text(file_content)
        storage_result = store_resume_file(reference_no, filename, file_content, uploaded.content_type)

        UPLOAD_STORE[reference_no] = {
            'filename': filename,
            'content_type': uploaded.content_type,
            'content': file_content,
            'text_preview': text_preview,
            'storage_backend': storage_result['storage_backend'],
            'storage_uri': storage_result['storage_uri'],
            'storage_key': storage_result['storage_key'],
        }

        ResumeUploadRecord.objects.update_or_create(
            reference_no=reference_no,
            defaults={
                'original_filename': filename,
                'content_type': uploaded.content_type or '',
                'storage_backend': storage_result['storage_backend'],
                'storage_key': storage_result['storage_key'],
                'object_size': len(file_content),
                'object_etag': '',
                'status': 'uploaded',
            },
        )

        return Response(
            {
                'reference_no': reference_no,
                'filename': filename,
                'storage_backend': storage_result['storage_backend'],
                'storage_uri': storage_result['storage_uri'],
                'status': 'uploaded',
                'message': 'Resume accepted for parsing.',
            },
            status=status.HTTP_201_CREATED,
        )


class ResumeUploadUrlView(APIView):
    permission_classes = [AllowAny]

    def post(self, request):
        filename = str(request.data.get('filename', '')).strip()
        content_type = str(request.data.get('content_type', '')).strip()

        if not _valid_resume_filename(filename):
            return Response(
                {'detail': _unsupported_file_detail(filename)},
                status=status.HTTP_400_BAD_REQUEST,
            )

        if not is_s3_configured():
            return Response(
                {
                    'detail': 'S3 is not configured. Use direct /resumes/upload fallback for local storage.',
                    'storage_backend': 'local',
                },
                status=status.HTTP_400_BAD_REQUEST,
            )

        reference_no = f"REF-{str(uuid.uuid4())[:8].upper()}"
        presigned = generate_presigned_upload_url(reference_no, filename, content_type or None)

        ResumeUploadRecord.objects.update_or_create(
            reference_no=reference_no,
            defaults={
                'original_filename': filename,
                'content_type': content_type,
                'storage_backend': 's3',
                'storage_key': str(presigned['storage_key']),
                'status': 'url_issued',
            },
        )

        return Response(
            {
                'reference_no': reference_no,
                'upload_url': presigned['upload_url'],
                's3_key': presigned['storage_key'],
                'storage_uri': presigned['storage_uri'],
                'expires_in': presigned['expires_in'],
                'storage_backend': 's3',
                'message': 'Pre-signed upload URL generated.',
            },
            status=status.HTTP_200_OK,
        )


class ResumeRegisterUploadView(APIView):
    permission_classes = [AllowAny]

    def post(self, request):
        reference_no = str(request.data.get('reference_no', '')).strip()
        s3_key = str(request.data.get('s3_key', '')).strip()
        filename = str(request.data.get('filename', '')).strip()
        content_type = str(request.data.get('content_type', '')).strip()

        if not reference_no or not s3_key:
            return Response(
                {'detail': 'reference_no and s3_key are required.'},
                status=status.HTTP_400_BAD_REQUEST,
            )

        if not filename:
            return Response(
                {'detail': 'filename is required.'},
                status=status.HTTP_400_BAD_REQUEST,
            )

        if not _valid_resume_filename(filename):
            return Response(
                {'detail': _unsupported_file_detail(filename)},
                status=status.HTTP_400_BAD_REQUEST,
            )

        object_size = None
        object_etag = ''
        if is_s3_configured():
            try:
                metadata = get_s3_object_metadata(s3_key)
                object_size = metadata.get('content_length')
                object_etag = str(metadata.get('etag') or '')
            except Exception:
                return Response(
                    {'detail': 'S3 object not found or not accessible for provided s3_key.'},
                    status=status.HTTP_400_BAD_REQUEST,
                )

        record, _ = ResumeUploadRecord.objects.update_or_create(
            reference_no=reference_no,
            defaults={
                'original_filename': filename,
                'content_type': content_type,
                'storage_backend': 's3',
                'storage_key': s3_key,
                'object_size': object_size,
                'object_etag': object_etag,
                'status': 'uploaded',
            },
        )

        UPLOAD_STORE[reference_no] = {
            'filename': record.original_filename,
            'content_type': record.content_type,
            'storage_backend': 's3',
            'storage_key': s3_key,
        }

        auto_parse = request.data.get('auto_parse', True)
        parsed_payload = None
        if bool(auto_parse):
            try:
                parsed_payload = _parse_and_persist_record(record)
            except Exception:
                record.status = 'failed'
                record.save(update_fields=['status', 'updated_at'])
                return Response(
                    {
                        'detail': 'Upload was registered but parsing failed.',
                        'reference_no': reference_no,
                    },
                    status=status.HTTP_500_INTERNAL_SERVER_ERROR,
                )

        return Response(
            {
                'reference_no': reference_no,
                'filename': record.original_filename,
                's3_key': s3_key,
                'object_size': record.object_size,
                'status': record.status,
                'message': 'Upload registered and parsed successfully.' if parsed_payload else 'Upload registered. Analysis can start.',
                'parsed_json': parsed_payload,
            },
            status=status.HTTP_200_OK,
        )


class ResumeParseView(APIView):
    permission_classes = [AllowAny]

    def post(self, request, reference_no: str):
        uploaded = UPLOAD_STORE.get(reference_no)
        record = ResumeUploadRecord.objects.filter(reference_no=reference_no).first()

        if not uploaded and not record:
            return Response(
                {
                    'detail': 'Unknown reference number. Upload resume before parsing.',
                    'reference_no': reference_no,
                },
                status=status.HTTP_404_NOT_FOUND,
            )

        if not record:
            record = ResumeUploadRecord.objects.filter(reference_no=reference_no).first()

        if not record or not record.storage_key:
            return Response(
                {
                    'detail': 'Upload exists but parsing payload is unavailable.',
                    'reference_no': reference_no,
                },
                status=status.HTTP_400_BAD_REQUEST,
            )

        try:
            parsed_payload = _parse_and_persist_record(record)
        except Exception:
            record.status = 'failed'
            record.save(update_fields=['status', 'updated_at'])
            return Response(
                {
                    'detail': 'Unable to load or parse uploaded file from storage.',
                    'reference_no': reference_no,
                },
                status=status.HTTP_500_INTERNAL_SERVER_ERROR,
            )

        return Response(parsed_payload, status=status.HTTP_200_OK)
